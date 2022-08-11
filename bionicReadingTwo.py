from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from tkinter import *
from logging import root
from tkinter import Label, Tk
from tkinter import filedialog
from tkinter import ttk
import string

def editWord(bionicText, blockPara, lastLetter):
    textLength = len(bionicText)
    if bionicText[0:textLength-1].isnumeric():
        blockPara.add_run(bionicText).bold = True
    elif textLength == 1:
        blockPara.add_run(bionicText).bold = True
    elif textLength == 4:
        blockPara.add_run(bionicText[0:2]).bold = True
        blockPara.add_run(bionicText[2:])
    elif textLength < 4:
        blockPara.add_run(bionicText).bold = True
    else:
        firstHalfNumber = ((textLength * 3)//5)
        if(bionicText[textLength-1] in string.punctuation):
            firstHalfNumber = ((textLength-1 * 3)//5)            
        blockPara.add_run(bionicText[0:int(firstHalfNumber)]).bold = True
        blockPara.add_run(bionicText[int(firstHalfNumber):])      
    if lastLetter == False:
        blockPara.add_run(" ")


def bionicProcess(blockText, blockPara):
    txtPosition = 1
    for bionictxt in blockText.split():
        if len(blockText.split()) == txtPosition:
            editWord(bionicText=bionictxt, blockPara= blockPara,lastLetter=True)
        else:
            editWord(bionicText=bionictxt, blockPara= blockPara,lastLetter=False)
        txtPosition = txtPosition + 1

        

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def recursive_iter_block_tems(parent):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                blockText = block.text
                if(block.text != ""):
                    block.clear()
                    bionicProcess(blockText=blockText,blockPara=block)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cells in row.cells:
                        recursive_iter_block_tems(cells)
def process_Document(fileToProcess):
    documentToProcess = Document(fileToProcess)
    recursive_iter_block_tems(documentToProcess)
    lastPosition = fileToProcess.rfind(".") 
    newFileLocation = fileToProcess[0:lastPosition] + r'-bionicReading.docx'
    documentToProcess.save(newFileLocation.replace("/", "\\"))
def select_file():
    filetypes = ((
        'document files', '*.docx'
    ),)
    selectedFile = filedialog.askopenfilename(title='Open a file', initialdir='/',
    filetypes=filetypes
    )
    process_Document(selectedFile)
def main():
    open_button = ttk.Button(root, text='Select a Document File', command=select_file)
    open_button.pack()

if __name__ == '__main__':
    root = Tk()
    root.title('Bionic Reading Converter')
    root.resizable(True, True)
    main()
    root.mainloop()